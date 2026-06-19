"""
APEX Pulse — Markdown to Word (.docx) converter
Converts all project .md files to branded Word documents in docs/word/
"""

import re
import os
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import docx.opc.constants
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

# ── Brand colours (for print: deep navy on white) ────────────────────────────
BRAND_PRIMARY   = RGBColor(0x18, 0x5F, 0xA5)   # #185FA5
BRAND_ACCENT    = RGBColor(0x37, 0x8A, 0xDD)   # #378ADD
BRAND_DARK      = RGBColor(0x07, 0x1E, 0x3D)   # #071E3D
BRAND_MUTED     = RGBColor(0x55, 0x6B, 0x82)   # muted blue-grey
COLOR_GREEN     = RGBColor(0x22, 0xA0, 0x6B)
COLOR_RED       = RGBColor(0xE2, 0x4B, 0x4A)
COLOR_AMBER     = RGBColor(0xC9, 0x8A, 0x20)
COLOR_GREY      = RGBColor(0x60, 0x60, 0x60)
COLOR_CODE_BG   = RGBColor(0xF4, 0xF6, 0xF9)   # light grey for code

FONT_BODY  = "Calibri"
FONT_MONO  = "Courier New"
FONT_HEAD  = "Calibri"


# ── Document setup ────────────────────────────────────────────────────────────

def new_document() -> Document:
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)
    # Normal style
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BRAND_DARK
    normal.paragraph_format.space_after = Pt(6)
    return doc


def set_heading_style(para, level: int):
    sizes = {1: 20, 2: 16, 3: 13, 4: 11.5}
    run = para.runs[0] if para.runs else para.add_run()
    run.font.name  = FONT_HEAD
    run.font.size  = Pt(sizes.get(level, 11))
    run.font.bold  = True
    run.font.color.rgb = BRAND_PRIMARY if level <= 2 else BRAND_ACCENT
    para.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    para.paragraph_format.space_after  = Pt(4)
    para.paragraph_format.keep_with_next = True


# ── Inline formatting ─────────────────────────────────────────────────────────

def add_inline(para, text: str):
    """Parse inline markdown (bold, italic, inline code, links) and add runs."""
    # Pattern order matters — code first (no nesting inside code)
    pattern = re.compile(
        r'`([^`]+)`'                        # `code`
        r'|\*\*\*(.+?)\*\*\*'              # ***bold italic***
        r'|\*\*(.+?)\*\*'                  # **bold**
        r'|__(.+?)__'                      # __bold__
        r'|\*(.+?)\*'                      # *italic*
        r'|_(.+?)_'                        # _italic_
        r'|\[([^\]]+)\]\(([^)]+)\)'        # [text](url)
        r'|~~(.+?)~~'                      # ~~strikethrough~~
    )
    pos = 0
    for m in pattern.finditer(text):
        # Plain text before this match
        if m.start() > pos:
            run = para.add_run(text[pos:m.start()])
            run.font.name = FONT_BODY
        g = m.groups()
        if g[0] is not None:   # `code`
            run = para.add_run(g[0])
            run.font.name  = FONT_MONO
            run.font.size  = Pt(9.5)
            run.font.color.rgb = BRAND_PRIMARY
        elif g[1] is not None: # ***bold italic***
            run = para.add_run(g[1])
            run.bold = True; run.italic = True
        elif g[2] is not None: # **bold**
            run = para.add_run(g[2])
            run.bold = True
            run.font.name = FONT_BODY
        elif g[3] is not None: # __bold__
            run = para.add_run(g[3])
            run.bold = True
        elif g[4] is not None: # *italic*
            run = para.add_run(g[4])
            run.italic = True
        elif g[5] is not None: # _italic_
            run = para.add_run(g[5])
            run.italic = True
        elif g[6] is not None: # [text](url)
            run = para.add_run(g[6])
            run.font.color.rgb = BRAND_ACCENT
            run.underline = True
        elif g[8] is not None: # ~~strikethrough~~
            run = para.add_run(g[8])
            run.font.strike = True
        pos = m.end()
    # Remaining plain text
    if pos < len(text):
        run = para.add_run(text[pos:])
        run.font.name = FONT_BODY


# ── Horizontal rule ───────────────────────────────────────────────────────────

def add_horizontal_rule(doc: Document):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(6)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "185FA5")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── Code block ────────────────────────────────────────────────────────────────

def add_code_block(doc: Document, lines: list[str], lang: str = ""):
    for i, line in enumerate(lines):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0) if i > 0 else Pt(6)
        para.paragraph_format.space_after  = Pt(0) if i < len(lines) - 1 else Pt(6)
        para.paragraph_format.left_indent  = Cm(0.5)
        # Shade the paragraph background
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F4F6F9")
        pPr.append(shd)
        run = para.add_run(line if line else " ")
        run.font.name  = FONT_MONO
        run.font.size  = Pt(8.5)
        run.font.color.rgb = BRAND_DARK


# ── Blockquote ────────────────────────────────────────────────────────────────

def add_blockquote(doc: Document, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Cm(1.0)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    # Left border
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "378ADD")
    pBdr.append(left)
    pPr.append(pBdr)
    run = para.add_run(text)
    run.font.italic    = True
    run.font.color.rgb = BRAND_MUTED
    run.font.name      = FONT_BODY


# ── Table ─────────────────────────────────────────────────────────────────────

def add_table(doc: Document, header: list[str], rows: list[list[str]]):
    n_cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, cell_text in enumerate(header):
        cell = hdr_row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(cell_text.strip())
        run.bold = True
        run.font.name = FONT_BODY
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Blue background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "185FA5")
        tcPr.append(shd)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        fill = "FFFFFF" if r_idx % 2 == 0 else "EEF4FB"
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= n_cols:
                break
            cell = row.cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            add_inline(para, cell_text.strip())
            for run in para.runs:
                run.font.size = Pt(9.5)
                run.font.name = FONT_BODY
            # Alternating row shading
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)

    doc.add_paragraph()  # spacing after table


# ── Cover page ────────────────────────────────────────────────────────────────

def add_cover_page(doc: Document, title: str, subtitle: str, filename: str):
    # Logo / brand mark area
    brand = doc.add_paragraph()
    brand.paragraph_format.space_before = Cm(2)
    brand.paragraph_format.space_after  = Pt(2)
    brand.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = brand.add_run("APEX Pulse")
    run.font.name  = FONT_HEAD
    run.font.size  = Pt(28)
    run.font.bold  = True
    run.font.color.rgb = BRAND_PRIMARY

    tagline = doc.add_paragraph("AI FinOps Intelligence Layer")
    tagline.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tagline.paragraph_format.space_after = Cm(1)
    run = tagline.runs[0]
    run.font.name  = FONT_BODY
    run.font.size  = Pt(11)
    run.font.color.rgb = BRAND_ACCENT

    # Separator line
    add_horizontal_rule(doc)

    # Document title
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.paragraph_format.space_before = Cm(1)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = t.add_run(title)
    run.font.name  = FONT_HEAD
    run.font.size  = Pt(22)
    run.font.bold  = True
    run.font.color.rgb = BRAND_DARK

    if subtitle:
        s = doc.add_paragraph(subtitle)
        s.paragraph_format.space_after = Cm(0.5)
        run = s.runs[0]
        run.font.name  = FONT_BODY
        run.font.size  = Pt(11)
        run.font.color.rgb = BRAND_MUTED

    # Metadata
    doc.add_paragraph()
    meta_lines = [
        f"Version: 10.0.0 · Sprint 10",
        f"Generated: {datetime.now().strftime('%d %B %Y')}",
        f"Source: {filename}",
        "Classification: Internal",
    ]
    for line in meta_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.runs[0]
        run.font.name  = FONT_MONO
        run.font.size  = Pt(9)
        run.font.color.rgb = BRAND_MUTED

    # Page break after cover
    doc.add_page_break()


# ── Header / Footer ───────────────────────────────────────────────────────────

def add_header_footer(doc: Document, doc_title: str):
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run(f"APEX Pulse  ·  {doc_title}")
        run.font.name  = FONT_BODY
        run.font.size  = Pt(8.5)
        run.font.color.rgb = BRAND_MUTED
        run.italic = True
        # Header bottom border
        pPr = hp._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "4")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "378ADD")
        pBdr.append(bot)
        pPr.append(pBdr)

        # Footer with page number
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_left = fp.add_run("APEX Pulse  ·  Confidential  ·  ")
        run_left.font.name  = FONT_BODY
        run_left.font.size  = Pt(8)
        run_left.font.color.rgb = BRAND_MUTED
        # Page number field
        fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText"); instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
        run_pg = fp.add_run()
        run_pg.font.size = Pt(8)
        run_pg.font.color.rgb = BRAND_MUTED
        run_pg._r.append(fldChar1); run_pg._r.append(instrText); run_pg._r.append(fldChar2)


# ── Main parser ───────────────────────────────────────────────────────────────

def parse_md_to_doc(md_path: Path, out_path: Path, doc_title: str, subtitle: str = ""):
    print(f"  Converting: {md_path.name} → {out_path.name}")
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = new_document()
    add_cover_page(doc, doc_title, subtitle, md_path.name)
    add_header_footer(doc, doc_title)

    i = 0
    in_code  = False
    code_buf = []
    code_lang = ""
    table_header = None
    table_rows   = []

    def flush_table():
        nonlocal table_header, table_rows
        if table_header:
            add_table(doc, table_header, table_rows)
        table_header = None
        table_rows   = []

    while i < len(lines):
        line = lines[i]

        # ── Code block toggle ──────────────────────────────────────────────
        if line.startswith("```"):
            if not in_code:
                flush_table()
                in_code   = True
                code_lang = line[3:].strip()
                code_buf  = []
            else:
                add_code_block(doc, code_buf, code_lang)
                in_code  = False
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ── Blank line ─────────────────────────────────────────────────────
        if not line.strip():
            flush_table()
            i += 1
            continue

        # ── Horizontal rule ────────────────────────────────────────────────
        if re.match(r'^[-*_]{3,}\s*$', line):
            flush_table()
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Headings ───────────────────────────────────────────────────────
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            flush_table()
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            style_name = f"Heading {min(level, 4)}"
            para = doc.add_paragraph(style=style_name)
            para.clear()
            add_inline(para, heading_text)
            set_heading_style(para, level)
            i += 1
            continue

        # ── Table ──────────────────────────────────────────────────────────
        if line.strip().startswith("|") and "|" in line:
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            # Check next line for separator
            if i + 1 < len(lines) and re.match(r'^\|[-| :]+\|?\s*$', lines[i + 1]):
                flush_table()
                table_header = cells
                table_rows   = []
                i += 2  # skip separator row
                continue
            elif table_header is not None:
                table_rows.append(cells)
                i += 1
                continue
            else:
                # Unknown table continuation — treat as text
                pass

        # ── Blockquote ─────────────────────────────────────────────────────
        if line.startswith(">"):
            flush_table()
            bq_text = re.sub(r'^>\s?', '', line)
            add_blockquote(doc, bq_text)
            i += 1
            continue

        # ── Unordered list ─────────────────────────────────────────────────
        if re.match(r'^(\s*)[-*+]\s+', line):
            flush_table()
            indent = len(re.match(r'^(\s*)', line).group(1))
            item_text = re.sub(r'^\s*[-*+]\s+', '', line)
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            para = doc.add_paragraph(style=style)
            para.clear()
            add_inline(para, item_text)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            for run in para.runs:
                run.font.name = FONT_BODY
                run.font.size = Pt(10.5)
            i += 1
            continue

        # ── Ordered list ───────────────────────────────────────────────────
        if re.match(r'^\d+\.\s+', line):
            flush_table()
            item_text = re.sub(r'^\d+\.\s+', '', line)
            para = doc.add_paragraph(style="List Number")
            para.clear()
            add_inline(para, item_text)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after  = Pt(1)
            for run in para.runs:
                run.font.name = FONT_BODY
                run.font.size = Pt(10.5)
            i += 1
            continue

        # ── Normal paragraph ───────────────────────────────────────────────
        flush_table()
        para = doc.add_paragraph()
        add_inline(para, line.strip())
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(6)
        for run in para.runs:
            if not run.font.name or run.font.name == "Calibri":
                run.font.name = FONT_BODY
            if not run.font.size:
                run.font.size = Pt(10.5)
        i += 1

    flush_table()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"    ✓  Saved: {out_path}")


# ── File manifest ─────────────────────────────────────────────────────────────

ROOT = Path(__file__).parent.parent
OUT  = ROOT / "docs" / "word"

MANIFEST = [
    (
        ROOT / "README.md",
        OUT  / "APEX_Pulse_Overview.docx",
        "APEX Pulse — Platform Overview",
        "Multi-provider AI cost observability · Sprint 10 · v10.0.0",
    ),
    (
        ROOT / "docs" / "SETUP.md",
        OUT  / "APEX_Pulse_Setup_Guide.docx",
        "APEX Pulse — Setup Guide",
        "Development, Docker, PostgreSQL, Nginx enterprise deployment",
    ),
    (
        ROOT / "docs" / "SECURITY.md",
        OUT  / "APEX_Pulse_Security_Reference.docx",
        "APEX Pulse — Security Reference",
        "Authentication, RBAC, TLS, audit trail, secrets management",
    ),
    (
        ROOT / "docs" / "DEPLOYMENT.md",
        OUT  / "APEX_Pulse_Deployment_Guide.docx",
        "APEX Pulse — Deployment Guide",
        "Vercel, Railway, Render, Docker, Kubernetes",
    ),
    (
        ROOT / "docs" / "INTEGRATIONS.md",
        OUT  / "APEX_Pulse_Integrations_Guide.docx",
        "APEX Pulse — Integrations Guide",
        "Anthropic, OpenAI, Google, Azure, Bedrock, Slack, webhooks",
    ),
    (
        ROOT / "docs" / "API.md",
        OUT  / "APEX_Pulse_API_Reference.docx",
        "APEX Pulse — API Reference",
        "REST endpoints, authentication, rate limiting, error codes",
    ),
    (
        ROOT / "docs" / "BACKLOG.md",
        OUT  / "APEX_Pulse_Backlog_S11_S20.docx",
        "APEX Pulse — Product Backlog & Sprint Plans",
        "Future sprints S11–S20 · Integration marketplace · SSO · Multi-tenant",
    ),
]


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print(f"\n{'='*60}")
    print("  APEX Pulse — Markdown → Word Converter")
    print(f"  Output: {OUT}")
    print(f"{'='*60}\n")

    success = 0
    for src, dst, title, subtitle in MANIFEST:
        if not src.exists():
            print(f"  SKIP (not found): {src.name}")
            continue
        try:
            parse_md_to_doc(src, dst, title, subtitle)
            success += 1
        except Exception as e:
            print(f"  ERROR converting {src.name}: {e}")
            import traceback; traceback.print_exc()

    print(f"\n{'='*60}")
    print(f"  Done: {success}/{len(MANIFEST)} files converted to {OUT}")
    print(f"{'='*60}\n")
